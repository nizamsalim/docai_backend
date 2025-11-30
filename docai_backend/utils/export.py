from docx import Document
from markdown import markdown
from bs4 import BeautifulSoup
from ..models.project_model import Project
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
from io import BytesIO


class DocumentExporter:

    @staticmethod
    def __make_black(paragraph_or_run):
        """Force all runs inside the element to black."""
        if hasattr(paragraph_or_run, "runs"):  # paragraph
            for run in paragraph_or_run.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        else:  # run
            paragraph_or_run.font.color.rgb = RGBColor(0, 0, 0)

    @staticmethod
    def __add_runs_from_html(element, paragraph):
        """Recursively add inline-formatted elements as docx runs."""
        for node in element.children:
            if node.name is None:
                run = paragraph.add_run(node.string)
                run.font.color.rgb = RGBColor(0, 0, 0)

            elif node.name in ["strong", "b"]:
                run = paragraph.add_run(node.get_text())
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            elif node.name in ["em", "i"]:
                run = paragraph.add_run(node.get_text())
                run.italic = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            elif node.name == "code":
                run = paragraph.add_run(node.get_text())
                run.font.name = "Consolas"
                run.font.color.rgb = RGBColor(0, 0, 0)

    @staticmethod
    def __md_to_docx(md_text: str, document: Document):
        html = markdown(md_text)
        soup = BeautifulSoup(html, "html.parser")

        for element in soup.children:
            # Headings
            if element.name and element.name.startswith("h"):
                level = int(element.name[1])
                paragraph = document.add_heading(level=level)
                DocumentExporter.__add_runs_from_html(element, paragraph)

            # Paragraphs
            elif element.name == "p":
                paragraph = document.add_paragraph()
                DocumentExporter.__add_runs_from_html(element, paragraph)

            # Unordered list
            elif element.name == "ul":
                for li in element.find_all("li", recursive=False):
                    paragraph = document.add_paragraph(style="List Bullet")
                    DocumentExporter.__add_runs_from_html(li, paragraph)

            # Ordered list
            elif element.name == "ol":
                for li in element.find_all("li", recursive=False):
                    paragraph = document.add_paragraph(style="List Number")
                    DocumentExporter.__add_runs_from_html(li, paragraph)

    @staticmethod
    def generate_word(project: Project) -> BytesIO:
        doc = Document()

        buffer = BytesIO()

        # Title
        title_paragraph = doc.add_heading(project.title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        doc.add_paragraph()  # blank space under title

        # Each section
        for section in project.sections:
            # Optional: Section title (Heading 1)
            # heading = doc.add_heading(section.title, level=1)

            # Section content from markdown
            DocumentExporter.__md_to_docx(section.content, doc)

            # Add small spacing instead of a page break
            space = doc.add_paragraph()
            space.paragraph_format.space_after = Pt(18)

        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_ppt(project: Project):
        pass
